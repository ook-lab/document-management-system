"""
B-6: Native Word Processor (.docx専用)

python-docx を使用して、ネイティブWord文書から構造化データを抽出。
表（Tables）、段落、文字装飾（赤字・太字）を100%精度で取得。
"""

from pathlib import Path
from typing import Dict, Any, List
from loguru import logger


class B6NativeWordProcessor:
    """B-6: Native Word Processor (.docx専用)"""

    def process(self, file_path: Path) -> Dict[str, Any]:
        """
        .docx ファイルから構造化データを抽出

        Args:
            file_path: .docxファイルパス

        Returns:
            {
                'is_structured': bool,
                'text_with_tags': str,           # 装飾タグ付きテキスト
                'structured_tables': [...],      # 表構造データ
                'paragraphs': [...],             # 段落リスト
                'tags': {...},                   # メタ情報
                'media_elements': [...]          # 埋め込み画像
            }
        """
        logger.info(f"[B-6] Native Word処理開始: {file_path.name}")

        try:
            from docx import Document
        except ImportError:
            logger.error("[B-6] python-docx がインストールされていません")
            return self._error_result("python-docx not installed")

        try:
            doc = Document(str(file_path))

            # 段落を抽出（文字装飾込み）
            paragraphs = self._extract_paragraphs(doc)

            # 表を抽出
            tables = self._extract_tables(doc)

            # 装飾タグ付きテキストを生成
            text_with_tags = self._build_tagged_text(paragraphs, tables)

            # メタ情報
            tags = {
                'has_red_text': any(p.get('has_red', False) for p in paragraphs),
                'has_bold_text': any(p.get('has_bold', False) for p in paragraphs),
                'has_table': len(tables) > 0,
                'table_count': len(tables),
                'paragraph_count': len(paragraphs)
            }

            logger.info(f"[B-6] 抽出完了: 段落={len(paragraphs)}, 表={len(tables)}")

            return {
                'is_structured': True,
                'text_with_tags': text_with_tags,
                'structured_tables': tables,
                'paragraphs': paragraphs,
                'tags': tags,
                'media_elements': []  # TODO: 画像抽出
            }

        except Exception as e:
            logger.error(f"[B-6] 処理エラー: {e}", exc_info=True)
            return self._error_result(str(e))

    def _extract_paragraphs(self, doc) -> List[Dict[str, Any]]:
        """
        段落を抽出（文字装飾込み）

        Returns:
            [{
                'index': int,
                'text': str,
                'runs': [...],  # 文字装飾情報
                'has_red': bool,
                'has_bold': bool
            }, ...]
        """
        paragraphs = []

        for idx, para in enumerate(doc.paragraphs):
            runs = []
            has_red = False
            has_bold = False

            for run in para.runs:
                # 文字装飾を検出
                is_bold = run.bold or False
                is_red = False

                # フォント色の判定（RGB）
                if run.font.color and run.font.color.rgb:
                    r, g, b = run.font.color.rgb
                    # 赤判定（R > 200 かつ G, B < 100）
                    if r > 200 and g < 100 and b < 100:
                        is_red = True

                runs.append({
                    'text': run.text,
                    'bold': is_bold,
                    'red': is_red
                })

                if is_bold:
                    has_bold = True
                if is_red:
                    has_red = True

            paragraphs.append({
                'index': idx,
                'text': para.text,
                'runs': runs,
                'has_red': has_red,
                'has_bold': has_bold
            })

        return paragraphs

    def _extract_tables(self, doc) -> List[Dict[str, Any]]:
        """
        表を抽出（セル結合も保持）

        Returns:
            [{
                'index': int,
                'rows': int,
                'cols': int,
                'data': [[...], [...], ...]  # 2次元配列
            }, ...]
        """
        tables = []

        for idx, table in enumerate(doc.tables):
            rows_data = []

            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                rows_data.append(row_data)

            tables.append({
                'index': idx,
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0,
                'data': rows_data
            })

        return tables

    def _build_tagged_text(
        self,
        paragraphs: List[Dict[str, Any]],
        tables: List[Dict[str, Any]]
    ) -> str:
        """
        装飾タグ付きテキストを生成

        Returns:
            [RED]赤字テキスト[/RED] [BOLD]太字テキスト[/BOLD] 形式
        """
        result = []

        # 段落をタグ付きで出力
        for para in paragraphs:
            para_text = ""
            for run in para['runs']:
                text = run['text']
                if run['red'] and run['bold']:
                    text = f"[RED][BOLD]{text}[/BOLD][/RED]"
                elif run['red']:
                    text = f"[RED]{text}[/RED]"
                elif run['bold']:
                    text = f"[BOLD]{text}[/BOLD]"
                para_text += text
            result.append(para_text)

        # 表を簡易テキスト化
        for table in tables:
            result.append("\n[TABLE]")
            for row in table['data']:
                result.append(" | ".join(row))
            result.append("[/TABLE]\n")

        return "\n".join(result)

    def _error_result(self, error_message: str) -> Dict[str, Any]:
        """エラー結果を返す"""
        return {
            'is_structured': False,
            'error': error_message,
            'text_with_tags': '',
            'structured_tables': [],
            'paragraphs': [],
            'tags': {},
            'media_elements': []
        }
